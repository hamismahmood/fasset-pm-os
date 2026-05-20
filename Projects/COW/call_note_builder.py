"""
COW Call Note Builder
Usage: edit the DATA section at the bottom, then run.
Builds a formatted .docx and uploads it to Google Drive as a Google Doc.
"""

import json
from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TOKEN_PATH = '/Users/hamis.mahmood/.claude/.sheets_token.json'
DOCX_MIME  = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
GDOC_MIME  = 'application/vnd.google-apps.document'

# ── helpers ──────────────────────────────────────────────────────────────────

def hex_rgb(h):
    h = h.lstrip('#')
    return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

def cell_bg(cell, color):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'), color.lstrip('#')); pr.append(shd)

def cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc; pr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for side, v in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        if v:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), v.get('val','single'))
            b.set(qn('w:sz'),  str(v.get('sz',4)))
            b.set(qn('w:color'), v.get('color','000000'))
            borders.append(b)
    pr.append(borders)

def no_border(table):
    tbl = table._tbl; pr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for side in ['top','bottom','left','right','insideH','insideV']:
        b = OxmlElement(f'w:{side}'); b.set(qn('w:val'),'none'); borders.append(b)
    pr.append(borders)

def section_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    r.font.size = Pt(8); r.font.bold = True
    r.font.color.rgb = hex_rgb('#64748b'); r.font.name = 'Calibri'
    pr = p._p.get_or_add_pPr(); bdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'4'); bot.set(qn('w:color'),'e2e8f0')
    bdr.append(bot); pr.append(bdr)

TAG_COLORS = {
    'PRODUCT':          ('#ede9fe','#5b21b6'),
    'UX':               ('#fef3c7','#92400e'),
    'FOLLOW-UP NEEDED': ('#fce7f3','#9d174d'),
    'INSIGHT':          ('#e0f2fe','#0c4a6e'),
}

# ── builder ───────────────────────────────────────────────────────────────────

def build_call_note(output_path, name, meta, who_is, why_happy,
                    stat, stat_label, stat_note, feedback_items, followups):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5); s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(2);  s.right_margin  = Cm(2)

    # Header
    ht = doc.add_table(rows=1, cols=1); no_border(ht)
    hc = ht.cell(0,0); cell_bg(hc, '#0f172a')
    lp = hc.paragraphs[0]
    lp.paragraph_format.space_before = Pt(4); lp.paragraph_format.space_after = Pt(2)
    lr = lp.add_run('USER CALL NOTE')
    lr.font.size=Pt(7.5); lr.font.bold=True; lr.font.color.rgb=hex_rgb('#94a3b8'); lr.font.name='Calibri'
    np_ = hc.add_paragraph()
    np_.paragraph_format.space_before=Pt(2); np_.paragraph_format.space_after=Pt(8)
    nr = np_.add_run(name)
    nr.font.size=Pt(22); nr.font.bold=True; nr.font.color.rgb=hex_rgb('#ffffff'); nr.font.name='Calibri'
    mp = hc.add_paragraph(); mp.paragraph_format.space_after=Pt(10)
    for i,m in enumerate(meta):
        if i:
            sep = mp.add_run('   ·   '); sep.font.color.rgb=hex_rgb('#475569'); sep.font.size=Pt(9.5); sep.font.name='Calibri'
        mr = mp.add_run(m); mr.font.size=Pt(9.5); mr.font.color.rgb=hex_rgb('#94a3b8'); mr.font.name='Calibri'
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Two-column body
    bt = doc.add_table(rows=1, cols=2); no_border(bt)
    bt.columns[0].width=Inches(3.1); bt.columns[1].width=Inches(3.4)
    lc = bt.cell(0,0); rc = bt.cell(0,1)

    def col_label(cell, text):
        p = cell.paragraphs[0]
        p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(6)
        r = p.add_run(text); r.font.size=Pt(7.5); r.font.bold=True
        r.font.color.rgb=hex_rgb('#64748b'); r.font.name='Calibri'
        pr=p._p.get_or_add_pPr(); bdr=OxmlElement('w:pBdr')
        b=OxmlElement('w:bottom'); b.set(qn('w:val'),'single'); b.set(qn('w:sz'),'4'); b.set(qn('w:color'),'e2e8f0')
        bdr.append(b); pr.append(bdr)

    col_label(lc, 'WHO HE IS')
    wp = lc.add_paragraph(who_is); wp.paragraph_format.space_before=Pt(4)
    for r in wp.runs: r.font.size=Pt(10.5); r.font.color.rgb=hex_rgb('#374151'); r.font.name='Calibri'

    col_label(rc, "WHY HE'S HAPPY")
    hyp = rc.add_paragraph(why_happy)
    hyp.paragraph_format.space_before=Pt(4); hyp.paragraph_format.space_after=Pt(6)
    for r in hyp.runs: r.font.size=Pt(10.5); r.font.color.rgb=hex_rgb('#374151'); r.font.name='Calibri'

    st = rc.add_table(rows=1, cols=1); no_border(st)
    sc = st.cell(0,0); cell_bg(sc,'#f0fdf4'); cell_border(sc, left={'val':'single','sz':18,'color':'22c55e'})
    sp = sc.paragraphs[0]; sp.paragraph_format.space_before=Pt(4); sp.paragraph_format.space_after=Pt(2)
    sr = sp.add_run(stat); sr.font.size=Pt(20); sr.font.bold=True; sr.font.color.rgb=hex_rgb('#15803d'); sr.font.name='Calibri'
    slp = sc.add_paragraph(stat_label); slp.paragraph_format.space_after=Pt(4)
    for r in slp.runs: r.font.size=Pt(9); r.font.color.rgb=hex_rgb('#166534'); r.font.name='Calibri'
    if stat_note:
        snp = rc.add_paragraph(stat_note); snp.paragraph_format.space_before=Pt(6)
        for r in snp.runs: r.font.size=Pt(9.5); r.font.color.rgb=hex_rgb('#4b5563'); r.font.name='Calibri'; r.font.italic=True

    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Feedback
    section_label(doc, 'FEEDBACK')
    for i, item in enumerate(feedback_items):
        ft = doc.add_table(rows=1, cols=2); no_border(ft)
        ft.columns[0].width=Inches(0.35); ft.columns[1].width=Inches(6.15)
        nc = ft.cell(0,0); cell_bg(nc,'#0f172a')
        np2 = nc.paragraphs[0]; np2.alignment=WD_ALIGN_PARAGRAPH.CENTER; np2.paragraph_format.space_before=Pt(4)
        nr2 = np2.add_run(str(i+1)); nr2.font.size=Pt(9); nr2.font.bold=True; nr2.font.color.rgb=hex_rgb('#ffffff'); nr2.font.name='Calibri'
        cc = ft.cell(0,1)
        cell_border(cc, top={'val':'single','sz':4,'color':'e2e8f0'},
                        bottom={'val':'single','sz':4,'color':'e2e8f0'},
                        right={'val':'single','sz':4,'color':'e2e8f0'})
        tp = cc.paragraphs[0]; tp.paragraph_format.space_before=Pt(6); tp.paragraph_format.space_after=Pt(3)
        tr = tp.add_run(item['title']+'  '); tr.font.size=Pt(10.5); tr.font.bold=True; tr.font.color.rgb=hex_rgb('#1a1a1a'); tr.font.name='Calibri'
        for tag in item.get('tags',[]):
            bg,fg = TAG_COLORS.get(tag,('#e2e8f0','#374151'))
            tagr = tp.add_run(f' {tag} '); tagr.font.size=Pt(7.5); tagr.font.bold=True; tagr.font.color.rgb=hex_rgb(fg); tagr.font.name='Calibri'
        dp = cc.add_paragraph(item['desc']); dp.paragraph_format.space_before=Pt(2); dp.paragraph_format.space_after=Pt(8)
        for r in dp.runs: r.font.size=Pt(9.5); r.font.color.rgb=hex_rgb('#4b5563'); r.font.name='Calibri'
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Follow-ups
    section_label(doc, 'FOLLOW-UPS')
    for fu in followups:
        fp = doc.add_paragraph(); fp.paragraph_format.space_before=Pt(3); fp.paragraph_format.space_after=Pt(3); fp.paragraph_format.left_indent=Inches(0.1)
        cbr = fp.add_run('☐  '); cbr.font.size=Pt(10); cbr.font.color.rgb=hex_rgb('#94a3b8'); cbr.font.name='Calibri'
        tr2 = fp.add_run(fu); tr2.font.size=Pt(10); tr2.font.color.rgb=hex_rgb('#374151'); tr2.font.name='Calibri'

    # Footer
    doc.add_paragraph().paragraph_format.space_before = Pt(16)
    ft2 = doc.add_table(rows=1, cols=2); no_border(ft2)
    cell_bg(ft2.cell(0,0),'#f8fafc'); cell_bg(ft2.cell(0,1),'#f8fafc')
    lf = ft2.cell(0,0).paragraphs[0]; lfr = lf.add_run('Fasset — Internal')
    lfr.font.size=Pt(8); lfr.font.color.rgb=hex_rgb('#94a3b8'); lfr.font.name='Calibri'
    rf = ft2.cell(0,1).paragraphs[0]; rf.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    rfr = rf.add_run(f'Hamis Mahmood  ·  {meta[0]}')
    rfr.font.size=Pt(8); rfr.font.color.rgb=hex_rgb('#94a3b8'); rfr.font.name='Calibri'

    doc.save(output_path)
    print(f'Saved: {output_path}')

# ── uploader ──────────────────────────────────────────────────────────────────

def upload_to_drive(docx_path, title):
    with open(TOKEN_PATH) as f:
        td = json.load(f)
    creds = Credentials(token=td['token'], refresh_token=td['refresh_token'],
                        token_uri=td['token_uri'], client_id=td['client_id'],
                        client_secret=td['client_secret'], scopes=td['scopes'])
    if creds.expired:
        creds.refresh(Request())
    svc = build('drive','v3', credentials=creds)
    meta = {'name': title, 'mimeType': GDOC_MIME}
    media = MediaFileUpload(docx_path, mimetype=DOCX_MIME, resumable=False)
    result = svc.files().create(body=meta, media_body=media, fields='id,webViewLink').execute()
    print(f'Uploaded: {title}')
    print(f'  Link: {result["webViewLink"]}')
    return result

# ── DATA — edit this section for each new call note ───────────────────────────

if __name__ == '__main__':
    OUTPUT = '/tmp/call_note_new.docx'

    build_call_note(
        output_path = OUTPUT,
        name        = 'User Name Here',
        meta        = ['DD Mon YYYY', '~XX minutes', 'Country', 'Context'],
        who_is      = 'Who the user is and their primary use case.',
        why_happy   = 'What they value about Fasset.',
        stat        = 'X%',
        stat_label  = 'One-line label for the stat',
        stat_note   = 'Optional supporting sentence.',
        feedback_items = [
            {
                'title': 'Feedback item title',
                'tags':  ['PRODUCT'],   # PRODUCT | UX | FOLLOW-UP NEEDED | INSIGHT
                'desc':  'Description of the feedback.',
            },
        ],
        followups = [
            'Follow-up action item 1.',
            'Follow-up action item 2.',
        ]
    )

    upload_to_drive(OUTPUT, 'Call Note — User Name Here (DD Mon YYYY)')
